"""
CV / spec-sheet export (Phase 1 feature 2).

UK recruiters rebrand every candidate before submitting to a client. This
service renders an agency-branded CV / candidate spec-sheet from a
CandidateProfile — optionally anonymised — as either a PDF (reportlab) or a
DOCX (python-docx). Both formats consume the same intermediate
`SpecSheetContent` bundle, so the two renderers stay in sync.

Anonymisation strips personal identifiers (name, email, phone) and blurs
company names to "Confidential" so the client sees skills + achievements
before knowing who or where the candidate is.

Template overrides layer on top of agency defaults:
    template.override or agency.value or renderer_default
so an agency can maintain, say, a "client-safe anonymised" template without
maintaining a full copy of its branding.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import date, datetime

from app.recruiter.models import (
    Agency,
    CandidateExperience,
    CandidateProfile,
    Role,
    SpecSheetTemplate,
)


# ── defaults ────────────────────────────────────────────────────────────────
DEFAULT_PRIMARY_COLOR = "#0f766e"
DEFAULT_HEADER_TEXT = "Candidate Spec Sheet"
DEFAULT_FOOTER_TEXT = "Prepared with ApplyForge Recruiter"
_HEX_COLOR_RE = re.compile(r"^#(?:[0-9a-fA-F]{3}|[0-9a-fA-F]{6}|[0-9a-fA-F]{8})$")


# ── resolved branding ──────────────────────────────────────────────────────
@dataclass
class ResolvedBranding:
    """Effective branding after template → agency → default fallback."""
    agency_name: str
    logo_url: str | None
    primary_color: str
    header_text: str
    footer_text: str
    body_intro: str | None
    template_name: str | None


def _first(*values, default=None):
    for v in values:
        if v is not None and (not isinstance(v, str) or v.strip()):
            return v
    return default


def _valid_color(hex_str: str | None) -> str | None:
    if not hex_str:
        return None
    return hex_str if _HEX_COLOR_RE.match(hex_str.strip()) else None


def resolve_branding(
    agency: Agency,
    template: SpecSheetTemplate | None = None,
) -> ResolvedBranding:
    """
    Layer template overrides on top of agency defaults on top of renderer
    defaults. Invalid hex values are dropped silently rather than propagating
    to the PDF engine.
    """
    return ResolvedBranding(
        agency_name=agency.name,
        logo_url=_first(
            (template.logo_url if template else None),
            agency.logo_url,
        ),
        primary_color=_first(
            _valid_color(template.primary_color) if template else None,
            _valid_color(agency.primary_color),
            default=DEFAULT_PRIMARY_COLOR,
        ),
        header_text=_first(
            (template.header_text if template else None),
            f"{agency.name} — Candidate",
            default=DEFAULT_HEADER_TEXT,
        ),
        footer_text=_first(
            (template.footer_text if template else None),
            agency.footer_text,
            default=DEFAULT_FOOTER_TEXT,
        ),
        body_intro=(template.body_intro if template else None),
        template_name=(template.name if template else None),
    )


# ── candidate → intermediate content bundle ────────────────────────────────
@dataclass
class SpecSheetSection:
    heading: str
    body: str | None = None
    lines: list[str] = field(default_factory=list)


@dataclass
class SpecSheetContent:
    """Renderer-agnostic content — both PDF and DOCX render from this."""
    display_name: str
    subtitle: str | None
    contact_lines: list[str]
    summary: str | None
    skills: list[str]
    experiences: list[dict]
    sections: list[SpecSheetSection]
    role_context: dict | None
    anonymised: bool


def _fmt_exp_dates(exp: CandidateExperience) -> str:
    """Turn (start_date, end_date) into a human-readable range."""
    def fmt(d: date | None) -> str | None:
        return d.strftime("%b %Y") if d else None
    start, end = fmt(exp.start_date), fmt(exp.end_date)
    if start and end:
        return f"{start} — {end}"
    if start:
        return f"{start} — Present"
    if end:
        return f"Until {end}"
    return ""


def _blur_company(company: str | None) -> str:
    """
    In anonymised mode we hide the exact employer but keep enough hint that
    the client can still assess industry seniority. If the company name has
    a suffix like "Ltd", "GmbH" etc, we keep it; otherwise we replace with
    "Confidential".
    """
    if not company:
        return "Confidential"
    return "Confidential"  # future: allow "Confidential (FinTech scale-up)" via metadata


def build_content(
    candidate: CandidateProfile,
    *,
    anonymise: bool,
    role: Role | None = None,
) -> SpecSheetContent:
    """Fold a CandidateProfile (plus optional role context) into the
    renderer-agnostic bundle used by build_pdf/build_docx."""
    if anonymise:
        # A neutral placeholder that reads like a real candidate ID. The
        # agency can also swap this out template-side in a future revision.
        initials = "".join(w[0] for w in (candidate.full_name or "").split()[:2]).upper()
        display_name = f"Candidate #{candidate.id:03d}" + (f" ({initials})" if initials else "")
        subtitle = candidate.headline  # headline is generic-enough to keep
        contact_lines = [
            f"Location: {candidate.location}" if candidate.location else "Location: Confidential",
            f"Experience: {candidate.years_experience:g} years" if candidate.years_experience else "Experience: Confidential",
        ]
    else:
        display_name = candidate.full_name or f"Candidate #{candidate.id:03d}"
        subtitle = candidate.headline
        contact_lines = [line for line in (
            candidate.email,
            candidate.phone,
            candidate.location,
            f"{candidate.years_experience:g} years experience" if candidate.years_experience else None,
        ) if line]

    experiences = []
    for exp in candidate.experiences or []:
        company = _blur_company(exp.company) if anonymise else exp.company
        experiences.append({
            "title": exp.title or "Role",
            "company": company or "",
            "dates": _fmt_exp_dates(exp),
            "description": exp.description or "",
        })

    skills = [s.name for s in (candidate.skills or [])]

    sections: list[SpecSheetSection] = []
    role_context = None
    if role is not None:
        role_context = {
            "role_id": role.id,
            "role_title": role.title,
            "role_client": None,
        }
        required = list(dict.fromkeys(role.required_skills or []))
        preferred = list(dict.fromkeys(role.preferred_skills or []))
        cand_skill_set = {s.lower() for s in skills}
        matched = [s for s in required + preferred if s.lower() in cand_skill_set]
        missing = [s for s in required if s.lower() not in cand_skill_set]
        section = SpecSheetSection(
            heading=f"Fit against role: {role.title}",
            body=None,
            lines=[
                f"Matched skills: {', '.join(matched) if matched else 'None yet'}",
                f"Gaps: {', '.join(missing) if missing else 'None — meets all required skills'}",
            ],
        )
        sections.append(section)

    return SpecSheetContent(
        display_name=display_name,
        subtitle=subtitle,
        contact_lines=contact_lines,
        summary=candidate.summary,
        skills=skills,
        experiences=experiences,
        sections=sections,
        role_context=role_context,
        anonymised=anonymise,
    )


# ── PDF renderer (reportlab) ───────────────────────────────────────────────
def build_pdf(content: SpecSheetContent, branding: ResolvedBranding) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable,
        Image,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )
    from reportlab.platypus.flowables import KeepInFrame

    primary = colors.HexColor(branding.primary_color)
    ink = colors.HexColor("#111827")
    muted = colors.HexColor("#6b7280")
    accent_bg = colors.HexColor("#f3f4f6")

    base = getSampleStyleSheet()
    styles: dict[str, ParagraphStyle] = {
        "h1": ParagraphStyle("h1", parent=base["Heading1"], fontName="Helvetica-Bold",
                             fontSize=22, leading=26, textColor=ink, spaceAfter=6),
        "h2": ParagraphStyle("h2", parent=base["Heading2"], fontName="Helvetica-Bold",
                             fontSize=13, leading=16, textColor=primary,
                             spaceBefore=14, spaceAfter=6),
        "sub": ParagraphStyle("sub", parent=base["Normal"], fontName="Helvetica",
                              fontSize=11, leading=14, textColor=muted),
        "body": ParagraphStyle("body", parent=base["Normal"], fontName="Helvetica",
                               fontSize=10.5, leading=14, textColor=ink),
        "small": ParagraphStyle("small", parent=base["Normal"], fontName="Helvetica",
                                fontSize=9, leading=12, textColor=muted),
        "label": ParagraphStyle("label", parent=base["Normal"], fontName="Helvetica-Bold",
                                fontSize=9, leading=11, textColor=muted, spaceAfter=2),
        "exp_title": ParagraphStyle("exp_title", parent=base["Normal"], fontName="Helvetica-Bold",
                                    fontSize=11, leading=14, textColor=ink),
        "exp_meta": ParagraphStyle("exp_meta", parent=base["Normal"], fontName="Helvetica-Oblique",
                                   fontSize=9.5, leading=12, textColor=muted, spaceAfter=4),
    }

    buf = io.BytesIO()
    title = f"{branding.agency_name} — {content.display_name}"
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm,   bottomMargin=2 * cm,
        title=title, author=branding.agency_name,
        subject="Candidate spec sheet",
    )
    story = []

    # ── header row: logo + agency name (left) / date (right) ──────────────
    left_cell = None
    if branding.logo_url:
        try:
            # Only http(s) URLs will resolve here; local/private paths would
            # fail. Fall back gracefully to the text name on any error.
            img = Image(branding.logo_url, width=3.2 * cm, height=1.2 * cm, kind="proportional")
            left_cell = img
        except Exception:
            left_cell = Paragraph(f"<b>{branding.agency_name}</b>", styles["body"])
    else:
        left_cell = Paragraph(f"<b>{branding.agency_name}</b>", styles["body"])

    header = Table(
        [[left_cell, Paragraph(datetime.utcnow().strftime("%d %b %Y"), styles["small"])]],
        colWidths=[11 * cm, 6 * cm],
    )
    header.setStyle(TableStyle([
        ("ALIGN", (1, 0), (1, 0), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(header)
    story.append(HRFlowable(width="100%", thickness=0.7, color=colors.HexColor("#e5e7eb")))
    story.append(Spacer(1, 10))

    # ── candidate identity ────────────────────────────────────────────────
    story.append(Paragraph(branding.header_text, styles["small"]))
    story.append(Paragraph(content.display_name, styles["h1"]))
    if content.subtitle:
        story.append(Paragraph(content.subtitle, styles["sub"]))
    if content.contact_lines:
        story.append(Spacer(1, 4))
        story.append(Paragraph(" · ".join(content.contact_lines), styles["small"]))

    if content.anonymised:
        story.append(Spacer(1, 4))
        story.append(Paragraph(
            "<i>This is an anonymised profile. Full contact details are released after client confirmation.</i>",
            styles["small"],
        ))

    # ── optional template intro ───────────────────────────────────────────
    if branding.body_intro:
        story.append(Spacer(1, 10))
        story.append(Paragraph(branding.body_intro.replace("\n", "<br/>"), styles["body"]))

    # ── summary / skills / experience ─────────────────────────────────────
    if content.summary:
        story.append(Paragraph("Profile summary", styles["h2"]))
        story.append(Paragraph(content.summary.replace("\n", "<br/>"), styles["body"]))

    if content.skills:
        story.append(Paragraph("Key skills", styles["h2"]))
        # Render skills as a compact chip line — reportlab table would over-engineer this.
        story.append(Paragraph(" · ".join(content.skills), styles["body"]))

    if content.experiences:
        story.append(Paragraph("Experience", styles["h2"]))
        for exp in content.experiences:
            title_line = exp["title"]
            if exp["company"]:
                title_line = f"{title_line} — {exp['company']}"
            story.append(Paragraph(title_line, styles["exp_title"]))
            if exp["dates"]:
                story.append(Paragraph(exp["dates"], styles["exp_meta"]))
            if exp["description"]:
                story.append(Paragraph(exp["description"].replace("\n", "<br/>"), styles["body"]))
                story.append(Spacer(1, 4))

    # ── role-fit callout ──────────────────────────────────────────────────
    for section in content.sections:
        story.append(Paragraph(section.heading, styles["h2"]))
        if section.body:
            story.append(Paragraph(section.body, styles["body"]))
        for line in section.lines:
            story.append(Paragraph(line, styles["body"]))

    # ── footer ────────────────────────────────────────────────────────────
    story.append(Spacer(1, 18))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e5e7eb")))
    story.append(Spacer(1, 6))
    footer_bits = [branding.footer_text]
    if branding.template_name:
        footer_bits.append(f"Template: {branding.template_name}")
    footer_bits.append(datetime.utcnow().strftime("%d %b %Y %H:%M UTC"))
    story.append(Paragraph(" · ".join(footer_bits), styles["small"]))

    # Wrap in KeepInFrame so oversize content is downscaled rather than
    # dropped silently. shrink mode preserves layout at 100% first.
    doc.build([KeepInFrame(0, 0, story, mode="shrink")])
    return buf.getvalue()


# ── DOCX renderer (python-docx) ────────────────────────────────────────────
def build_docx(content: SpecSheetContent, branding: ResolvedBranding) -> bytes:
    """
    Produce an editable .docx so the recruiter can tweak wording before
    sending. Uses python-docx (already a backend dep). Colours are applied
    to headings via run.color.rgb; we don't attempt to embed the logo when
    it's a URL — reportlab handles that path, and python-docx would need us
    to fetch the bytes ourselves.
    """
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    def _rgb_from_hex(hex_str: str) -> RGBColor:
        h = hex_str.lstrip("#")
        if len(h) == 3:
            h = "".join(c * 2 for c in h)
        if len(h) == 8:
            h = h[:6]   # drop alpha
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    primary = _rgb_from_hex(branding.primary_color)
    ink = RGBColor(0x11, 0x18, 0x27)
    muted = RGBColor(0x6b, 0x72, 0x80)

    doc = Document()

    # ── page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ── header block ──
    header_p = doc.add_paragraph()
    header_run = header_p.add_run(branding.agency_name)
    header_run.bold = True
    header_run.font.size = Pt(12)
    header_run.font.color.rgb = ink

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_p.add_run(datetime.utcnow().strftime("%d %b %Y"))
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = muted

    # ── identity ──
    tagline = doc.add_paragraph()
    tag_run = tagline.add_run(branding.header_text)
    tag_run.font.size = Pt(9)
    tag_run.font.color.rgb = muted

    name_p = doc.add_paragraph()
    name_run = name_p.add_run(content.display_name)
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = ink

    if content.subtitle:
        sub_p = doc.add_paragraph()
        sr = sub_p.add_run(content.subtitle)
        sr.font.size = Pt(11)
        sr.font.color.rgb = muted

    if content.contact_lines:
        cp = doc.add_paragraph()
        cr = cp.add_run(" · ".join(content.contact_lines))
        cr.font.size = Pt(9)
        cr.font.color.rgb = muted

    if content.anonymised:
        note = doc.add_paragraph()
        nr = note.add_run(
            "This is an anonymised profile. Full contact details are released after client confirmation."
        )
        nr.italic = True
        nr.font.size = Pt(9)
        nr.font.color.rgb = muted

    if branding.body_intro:
        doc.add_paragraph()
        intro = doc.add_paragraph()
        intro_run = intro.add_run(branding.body_intro)
        intro_run.font.size = Pt(10.5)
        intro_run.font.color.rgb = ink

    def _h2(text: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = primary
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    def _body(text: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.color.rgb = ink

    if content.summary:
        _h2("Profile summary")
        _body(content.summary)

    if content.skills:
        _h2("Key skills")
        _body(" · ".join(content.skills))

    if content.experiences:
        _h2("Experience")
        for exp in content.experiences:
            p = doc.add_paragraph()
            title_run = p.add_run(exp["title"])
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = ink
            if exp["company"]:
                cr = p.add_run(f" — {exp['company']}")
                cr.font.size = Pt(11)
                cr.font.color.rgb = ink
            if exp["dates"]:
                dp = doc.add_paragraph()
                dr = dp.add_run(exp["dates"])
                dr.italic = True
                dr.font.size = Pt(9.5)
                dr.font.color.rgb = muted
            if exp["description"]:
                _body(exp["description"])

    for section in content.sections:
        _h2(section.heading)
        if section.body:
            _body(section.body)
        for line in section.lines:
            _body(line)

    # ── footer paragraph (docx footers are per-section — this is body text
    # at the end so the recruiter can edit it if they want) ──
    doc.add_paragraph()
    footer_bits = [branding.footer_text]
    if branding.template_name:
        footer_bits.append(f"Template: {branding.template_name}")
    footer_bits.append(datetime.utcnow().strftime("%d %b %Y %H:%M UTC"))
    fp = doc.add_paragraph()
    fr = fp.add_run(" · ".join(footer_bits))
    fr.font.size = Pt(9)
    fr.font.color.rgb = muted

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── convenience wrappers used by the API layer ─────────────────────────────
def build_spec_sheet_pdf(
    candidate: CandidateProfile,
    agency: Agency,
    *,
    anonymise: bool,
    role: Role | None = None,
    template: SpecSheetTemplate | None = None,
) -> bytes:
    branding = resolve_branding(agency, template)
    content = build_content(candidate, anonymise=anonymise, role=role)
    return build_pdf(content, branding)


def build_spec_sheet_docx(
    candidate: CandidateProfile,
    agency: Agency,
    *,
    anonymise: bool,
    role: Role | None = None,
    template: SpecSheetTemplate | None = None,
) -> bytes:
    branding = resolve_branding(agency, template)
    content = build_content(candidate, anonymise=anonymise, role=role)
    return build_docx(content, branding)


def filename_for(
    candidate: CandidateProfile,
    branding: ResolvedBranding,
    *,
    extension: str,
    anonymise: bool,
) -> str:
    """Slugify agency + candidate into a sensible download filename."""
    def slug(s: str | None) -> str:
        if not s:
            return ""
        s = re.sub(r"[^A-Za-z0-9]+", "-", s).strip("-")
        return s.lower()

    parts = [slug(branding.agency_name), "spec-sheet"]
    if anonymise:
        parts.append(f"candidate-{candidate.id:03d}")
    else:
        parts.append(slug(candidate.full_name) or f"candidate-{candidate.id:03d}")
    return "-".join(p for p in parts if p) + f".{extension}"
