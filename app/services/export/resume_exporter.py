"""
Resume Export Service
=====================
Generates downloadable PDF and DOCX versions of the user's resume,
sourced from the normalised profile tables (same data the ATS preview
renders). Falls back to ParsedResumeData for any section that has no
DB rows yet (matching the frontend merge logic).

Libraries used:
  PDF  — reportlab  (already installed)
  DOCX — python-docx (already installed)
"""
from __future__ import annotations

import io
import json
from datetime import date, datetime
from typing import Any

import structlog
from fastapi import HTTPException
from sqlalchemy.orm import Session

from app.models.models import (
    ApplicationCustomization,
    Certification,
    Education,
    ParsedResumeData,
    Project,
    Skill,
    User,
    UserProfile,
    WorkExperience,
)

logger = structlog.get_logger(__name__)


# ── helpers ────────────────────────────────────────────────────────────────

def _fmt_date(d: date | datetime | str | None) -> str:
    if d is None:
        return "Present"
    if isinstance(d, str):
        try:
            d = datetime.fromisoformat(d)
        except ValueError:
            return d
    return d.strftime("%b %Y")


def _bullets(description: str | None, max_bullets: int = 6) -> list[str]:
    """Split a description text into bullet lines."""
    if not description:
        return []
    lines = [l.strip().lstrip("•–-").strip() for l in description.splitlines()]
    lines = [l for l in lines if l]
    if len(lines) > 1:
        return lines[:max_bullets]
    # single paragraph — split on ". "
    parts = [s.strip() for s in description.split(". ") if s.strip()]
    return parts[:max_bullets]


# ── data loader ────────────────────────────────────────────────────────────

class _ResumeData:
    """Loads all profile sections for a user, merging DB + parsed fallback.

    When *app_id* is supplied the per-application AI customizations are also
    applied (skills_add, experiences_update, projects_add) so the export
    matches exactly what the user sees in the on-screen preview.
    """

    def __init__(self, db: Session, user_id: int, user_email: str, app_id: int | None = None):
        profile = db.query(UserProfile).filter_by(user_id=user_id).first()
        parsed_row = (
            db.query(ParsedResumeData)
            .filter_by(user_id=user_id)
            .filter(ParsedResumeData.deleted_at.is_(None))
            .order_by(ParsedResumeData.created_at.desc())
            .first()
        )
        parsed: dict[str, Any] = {}
        if parsed_row and parsed_row.structured_json:
            try:
                raw = json.loads(parsed_row.structured_json)
                parsed = raw if isinstance(raw, dict) else {}
            except (json.JSONDecodeError, TypeError):
                parsed = {}

        # Basic info
        self.full_name: str = (
            (profile.full_name if profile else None)
            or parsed.get("full_name", "")
            or ""
        )
        self.headline: str = (
            (profile.headline if profile else None)
            or parsed.get("headline", "")
            or ""
        )
        self.summary: str = (
            (profile.summary if profile else None)
            or parsed.get("summary", "")
            or ""
        )
        self.location: str = (
            (profile.location if profile else None)
            or parsed.get("location", "")
            or ""
        )
        self.email: str = user_email

        # Per-application AI customizations (applied suggestions)
        customizations: dict[str, Any] = {}
        if app_id is not None:
            cust_row = (
                db.query(ApplicationCustomization)
                .filter_by(user_id=user_id, application_id=app_id)
                .first()
            )
            if cust_row and cust_row.customizations_json:
                try:
                    raw = json.loads(cust_row.customizations_json)
                    customizations = raw if isinstance(raw, dict) else {}
                except (json.JSONDecodeError, TypeError):
                    customizations = {}

        # Skills — merge DB + parsed (same dedup logic as frontend)
        db_skills = db.query(Skill).filter_by(user_id=user_id).all()
        seen: set[str] = set()
        skills: list[dict] = []
        for s in db_skills:
            key = (s.name or "").lower().strip()
            if key and key not in seen:
                seen.add(key)
                skills.append({"name": s.name, "level": s.level})
        for s in parsed.get("skills", []):
            item = {"name": s} if isinstance(s, str) else s
            key = (item.get("name", "") or "").lower().strip()
            if key and key not in seen:
                seen.add(key)
                skills.append(item)
        # AI-suggested skills for this specific job application
        for s in customizations.get("skills_add", []):
            key = (s.get("name", "") or "").lower().strip()
            if key and key not in seen:
                seen.add(key)
                skills.append({"name": s["name"], "level": s.get("level")})
        self.skills = skills

        # Experiences
        db_exps = (
            db.query(WorkExperience)
            .filter_by(user_id=user_id)
            .order_by(WorkExperience.start_date.desc())
            .all()
        )
        exp_overrides: dict[str, dict] = customizations.get("experiences_update", {})
        if db_exps:
            self.experiences = [
                {
                    "role": e.role,
                    "company": e.company,
                    "start_date": e.start_date,
                    "end_date": e.end_date,
                    # Apply AI-suggested bullet overrides for this job if present
                    "description": (
                        exp_overrides.get(str(e.id), {}).get("description")
                        or e.description
                    ),
                }
                for e in db_exps
            ]
        else:
            raw_exps = parsed.get("work_experience", [])
            self.experiences = sorted(
                raw_exps,
                key=lambda x: x.get("start_date") or "",
                reverse=True,
            )

        # Education — union merge: DB rows + parsed-resume (deduped by institution+degree)
        db_edus = (
            db.query(Education)
            .filter_by(user_id=user_id)
            .order_by(Education.start_date.desc())
            .all()
        )
        seen_edu: set[str] = set()
        educations: list[dict] = []
        for e in db_edus:
            key = f"{(e.institution or '').lower().strip()}::{(e.degree or '').lower().strip()}"
            if key not in seen_edu:
                seen_edu.add(key)
                educations.append({
                    "institution": e.institution,
                    "degree": e.degree,
                    "field_of_study": e.field_of_study,
                    "start_date": e.start_date,
                    "end_date": e.end_date,
                })
        for e in parsed.get("education", []):
            key = f"{(e.get('institution', '') or '').lower().strip()}::{(e.get('degree', '') or '').lower().strip()}"
            if key not in seen_edu:
                seen_edu.add(key)
                educations.append({
                    "institution": e.get("institution", ""),
                    "degree": e.get("degree", ""),
                    "field_of_study": e.get("field_of_study", ""),
                    "start_date": e.get("start_date"),
                    "end_date": e.get("end_date"),
                })
        self.educations = educations

        # Projects — union merge: DB rows + parsed-resume + AI-suggested
        # (mirrors the frontend ATSResumeTemplate union-merge logic)
        db_projs = db.query(Project).filter_by(user_id=user_id).all()
        seen_proj: set[str] = set()
        projects: list[dict] = []
        for p in db_projs:
            key = (p.name or "").lower().strip()
            if key and key not in seen_proj:
                seen_proj.add(key)
                projects.append({
                    "name": p.name,
                    "description": p.description,
                    "technologies": p.technologies,
                })
        for p in parsed.get("projects", []):
            key = (p.get("name", "") or "").lower().strip()
            if key and key not in seen_proj:
                seen_proj.add(key)
                projects.append({
                    "name": p.get("name", ""),
                    "description": p.get("description", ""),
                    "technologies": p.get("technologies", ""),
                })
        # Append AI-suggested projects (gap-fix or job-specific extras)
        for p in customizations.get("projects_add", []):
            key = (p.get("name", "") or "").lower().strip()
            if key and key not in seen_proj:
                seen_proj.add(key)
                projects.append({
                    "name": p.get("name", ""),
                    "description": p.get("description", ""),
                    "technologies": p.get("technologies", ""),
                })
        self.projects = projects

        # Certifications — union merge: DB rows + parsed-resume
        db_certs = db.query(Certification).filter_by(user_id=user_id).all()
        seen_cert: set[str] = set()
        certifications: list[dict] = []
        for c in db_certs:
            key = (c.name or "").lower().strip()
            if key and key not in seen_cert:
                seen_cert.add(key)
                certifications.append({"name": c.name, "issuer": c.issuer, "issue_date": c.issue_date})
        for c in parsed.get("certifications", []):
            key = (c.get("name", "") or "").lower().strip()
            if key and key not in seen_cert:
                seen_cert.add(key)
                certifications.append({
                    "name": c.get("name", ""),
                    "issuer": c.get("issuer", ""),
                    "issue_date": c.get("issue_date"),
                })
        self.certifications = certifications


# ── PDF export ─────────────────────────────────────────────────────────────

def _build_pdf(data: _ResumeData) -> bytes:
    """
    ATS-optimised single-column PDF resume.

    Design principles:
    - Single column only — multi-column trips up most ATS parsers
    - Standard section names (Professional Summary, Skills, Work Experience…)
    - Plain text bullets — no Unicode trickery
    - Dates right-aligned via a two-cell table (safest cross-renderer approach)
    - Skills rendered as comma-separated text (ATS reads prose, not grids)
    - No images, no text boxes, no headers/footers
    - Helvetica throughout (universal, renders identically on all platforms)
    """
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_JUSTIFY
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch, mm
    from reportlab.platypus import (
        HRFlowable,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
        KeepTogether,
    )

    # ── Dimensions ──────────────────────────────────────────────────────
    PAGE_W, PAGE_H = letter           # 8.5 × 11 in
    L_MAR = R_MAR = 0.65 * inch
    T_MAR = B_MAR = 0.55 * inch
    CONTENT_W = PAGE_W - L_MAR - R_MAR   # ≈ 7.2 in

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=L_MAR, rightMargin=R_MAR,
        topMargin=T_MAR, bottomMargin=B_MAR,
    )

    # ── Colour palette ───────────────────────────────────────────────────
    C_ACCENT  = colors.HexColor("#1D4ED8")   # strong blue  — section headers / company
    C_DARK    = colors.HexColor("#0F172A")   # near-black   — name, job titles
    C_BODY    = colors.HexColor("#1E293B")   # slate-900    — body text
    C_MID     = colors.HexColor("#475569")   # slate-600    — bullets, dates
    C_LIGHT   = colors.HexColor("#64748B")   # slate-500    — secondary labels
    C_RULE    = colors.HexColor("#CBD5E1")   # slate-300    — dividers
    C_ACCENT_RULE = colors.HexColor("#1D4ED8")

    # ── Style factory ────────────────────────────────────────────────────
    _base = getSampleStyleSheet()["Normal"]
    def S(name, **kw):
        return ParagraphStyle(name, parent=_base, **kw)

    # Header block
    sty_name = S("RName",
        fontName="Helvetica-Bold", fontSize=24, leading=29,
        textColor=C_DARK, spaceAfter=1)
    sty_headline = S("RHead",
        fontName="Helvetica", fontSize=11, leading=15,
        textColor=C_ACCENT, spaceAfter=3)
    sty_contact = S("RContact",
        fontName="Helvetica", fontSize=9, leading=13,
        textColor=C_LIGHT, spaceAfter=0)

    # Section headings — bold caps with a rule drawn below via HRFlowable
    sty_sec = S("RSec",
        fontName="Helvetica-Bold", fontSize=8.5, leading=11,
        textColor=C_ACCENT_RULE, spaceBefore=10, spaceAfter=2,
        letterSpacing=1.2)

    # Summary
    sty_summary = S("RSum",
        fontName="Helvetica", fontSize=9.5, leading=14.5,
        textColor=C_BODY, alignment=TA_JUSTIFY, spaceAfter=2)

    # Skills — plain paragraph, ATS-friendly
    sty_skills = S("RSkills",
        fontName="Helvetica", fontSize=9.5, leading=14.5,
        textColor=C_BODY, spaceAfter=2)

    # Experience
    sty_role = S("RRole",
        fontName="Helvetica-Bold", fontSize=10.5, leading=14,
        textColor=C_DARK, spaceAfter=0)
    sty_company = S("RCo",
        fontName="Helvetica-Oblique", fontSize=9.5, leading=13,
        textColor=C_ACCENT, spaceAfter=1)
    sty_date = S("RDate",
        fontName="Helvetica", fontSize=9, leading=12,
        textColor=C_MID, alignment=TA_RIGHT)
    sty_bullet = S("RBullet",
        fontName="Helvetica", fontSize=9.5, leading=14,
        textColor=C_BODY,
        leftIndent=14, firstLineIndent=0,
        spaceAfter=1.5, spaceBefore=0)

    # Projects
    sty_proj_name = S("RProjName",
        fontName="Helvetica-Bold", fontSize=10, leading=13,
        textColor=C_DARK, spaceAfter=1, spaceBefore=4)
    sty_proj_desc = S("RProjDesc",
        fontName="Helvetica", fontSize=9.5, leading=14,
        textColor=C_MID, spaceAfter=2, alignment=TA_JUSTIFY)

    # Education / Certs
    sty_inst = S("RInst",
        fontName="Helvetica-Bold", fontSize=10.5, leading=14,
        textColor=C_DARK, spaceAfter=0)
    sty_degree = S("RDeg",
        fontName="Helvetica-Oblique", fontSize=9.5, leading=13,
        textColor=C_MID, spaceAfter=4)
    sty_cert = S("RCert",
        fontName="Helvetica", fontSize=9.5, leading=14,
        textColor=C_BODY, spaceAfter=2)

    # ── Helpers ──────────────────────────────────────────────────────────
    def accent_rule():
        """Thin blue rule under section headings."""
        return HRFlowable(
            width="100%", thickness=1.2,
            color=C_ACCENT_RULE, spaceAfter=5, spaceBefore=0,
        )

    def light_rule():
        """Very faint separator between experience entries."""
        return HRFlowable(
            width="100%", thickness=0.4,
            color=C_RULE, spaceAfter=5, spaceBefore=4,
        )

    def section(title: str):
        return [Paragraph(title.upper(), sty_sec), accent_rule()]

    def exp_header(role: str, company: str, period: str) -> Table:
        """Role + company on left, date on right — kept in one line."""
        role_para    = Paragraph(role, sty_role)
        company_para = Paragraph(company, sty_company)
        date_para    = Paragraph(period, sty_date)
        tbl = Table(
            [[role_para, date_para],
             [company_para, ""]],
            colWidths=[CONTENT_W * 0.72, CONTENT_W * 0.28],
        )
        tbl.setStyle(TableStyle([
            ("VALIGN",        (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING",   (0, 0), (-1, -1), 0),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 0),
            ("TOPPADDING",    (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        return tbl

    def edu_header(institution: str, period: str) -> Table:
        tbl = Table(
            [[Paragraph(institution, sty_inst), Paragraph(period, sty_date)]],
            colWidths=[CONTENT_W * 0.72, CONTENT_W * 0.28],
        )
        tbl.setStyle(TableStyle([
            ("VALIGN",        (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING",   (0, 0), (-1, -1), 0),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 0),
            ("TOPPADDING",    (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        return tbl

    def bullet_para(text: str) -> Paragraph:
        """Bullet with a real en-dash leader — ATS reads the plain text fine."""
        return Paragraph(f"&#x2022;&#160;&#160;{text}", sty_bullet)

    # ── Build story ──────────────────────────────────────────────────────
    story = []

    # ── NAME ────────────────────────────────────────────────────────────
    story.append(Paragraph(data.full_name or "Resume", sty_name))
    if data.headline:
        story.append(Paragraph(data.headline, sty_headline))

    # Contact line — email · location (pipe-separated for ATS)
    contact_parts = []
    if data.email:
        contact_parts.append(data.email)
    if data.location:
        contact_parts.append(data.location)
    if contact_parts:
        story.append(Paragraph("  |  ".join(contact_parts), sty_contact))

    # Thick accent rule under the header block
    story.append(Spacer(1, 4))
    story.append(HRFlowable(
        width="100%", thickness=2, color=C_ACCENT,
        spaceAfter=8, spaceBefore=0,
    ))

    # ── PROFESSIONAL SUMMARY ─────────────────────────────────────────────
    if data.summary:
        story += section("Professional Summary")
        story.append(Paragraph(data.summary, sty_summary))
        story.append(Spacer(1, 4))

    # ── SKILLS ──────────────────────────────────────────────────────────
    if data.skills:
        story += section("Skills")
        # Comma-separated — every ATS parser can read this perfectly
        skill_names = [s["name"] for s in data.skills if s.get("name")]
        story.append(Paragraph(", ".join(skill_names), sty_skills))
        story.append(Spacer(1, 4))

    # ── WORK EXPERIENCE ──────────────────────────────────────────────────
    if data.experiences:
        story += section("Work Experience")
        for i, exp in enumerate(data.experiences):
            role    = exp.get("role") or ""
            company = exp.get("company") or ""
            start   = _fmt_date(exp.get("start_date"))
            end     = _fmt_date(exp.get("end_date"))
            period  = f"{start} \u2013 {end}"
            bullets = _bullets(exp.get("description"), 6)

            block = [exp_header(role, company, period)]
            for b in bullets:
                block.append(bullet_para(b))
            if i < len(data.experiences) - 1:
                block.append(light_rule())
            else:
                block.append(Spacer(1, 6))
            story.append(KeepTogether(block))

    # ── PROJECTS ────────────────────────────────────────────────────────
    if data.projects:
        story += section("Projects & Others")
        for proj in data.projects:
            name = proj.get("name") or ""
            tech = proj.get("technologies") or ""
            desc = proj.get("description") or ""
            tech_str = f" <font color='#1D4ED8'><i>| {tech}</i></font>" if tech else ""
            block = [Paragraph(f"{name}{tech_str}", sty_proj_name)]
            if desc:
                block.append(Paragraph(desc, sty_proj_desc))
            story.append(KeepTogether(block))
        story.append(Spacer(1, 4))

    # ── EDUCATION ────────────────────────────────────────────────────────
    if data.educations:
        story += section("Education")
        for edu in data.educations:
            institution = edu.get("institution") or ""
            degree      = edu.get("degree") or ""
            field       = edu.get("field_of_study") or ""
            start       = _fmt_date(edu.get("start_date"))
            end         = _fmt_date(edu.get("end_date"))
            degree_line = ", ".join(filter(None, [degree, field]))
            period      = f"{start} \u2013 {end}"
            block = [edu_header(institution, period)]
            if degree_line:
                block.append(Paragraph(degree_line, sty_degree))
            story.append(KeepTogether(block))
        story.append(Spacer(1, 4))

    # ── CERTIFICATIONS ───────────────────────────────────────────────────
    if data.certifications:
        story += section("Certifications")
        for cert in data.certifications:
            name   = cert.get("name") or ""
            issuer = cert.get("issuer") or ""
            issued = _fmt_date(cert.get("issue_date"))
            line_parts = [f"<b>{name}</b>"]
            if issuer:
                line_parts.append(issuer)
            if cert.get("issue_date"):
                line_parts.append(issued)
            story.append(Paragraph(" — ".join(line_parts), sty_cert))

    doc.build(story)
    return buf.getvalue()


# ── DOCX export ────────────────────────────────────────────────────────────

def _build_docx(data: _ResumeData) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml import OxmlElement

    ACCENT_RGB = RGBColor(0x1e, 0x40, 0xAF)   # blue-800
    DARK_RGB   = RGBColor(0x11, 0x18, 0x27)    # gray-900
    MID_RGB    = RGBColor(0x37, 0x41, 0x51)    # gray-700
    LIGHT_RGB  = RGBColor(0x6b, 0x72, 0x80)    # gray-500

    doc = Document()

    # ── page margins ────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    def add_run(para, text: str, bold=False, italic=False,
                color: RGBColor | None = None, size_pt: float = 10):
        run = para.add_run(text)
        run.bold   = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        run.font.size = Pt(size_pt)
        return run

    def section_heading(title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = ACCENT_RGB
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1e40af")
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    # ── name + headline ─────────────────────────────────────────────────
    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_after = Pt(1)
    add_run(name_p, data.full_name or "Resume", bold=True, color=DARK_RGB, size_pt=22)

    if data.headline:
        h_p = doc.add_paragraph()
        h_p.paragraph_format.space_after = Pt(2)
        add_run(h_p, data.headline, color=ACCENT_RGB, size_pt=11)

    contact_parts = [p for p in [data.email, data.location] if p]
    if contact_parts:
        c_p = doc.add_paragraph()
        c_p.paragraph_format.space_after = Pt(6)
        add_run(c_p, "  ·  ".join(contact_parts), color=LIGHT_RGB, size_pt=9)

    # ── summary ─────────────────────────────────────────────────────────
    if data.summary:
        section_heading("Professional Summary")
        p = doc.add_paragraph(data.summary)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = MID_RGB

    # ── skills ──────────────────────────────────────────────────────────
    if data.skills:
        section_heading("Skills")
        skill_names = [
            f"{s['name']} ({s['level']})" if s.get("level") else s["name"]
            for s in data.skills
        ]
        p = doc.add_paragraph("  ·  ".join(skill_names))
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = MID_RGB

    # ── experience ──────────────────────────────────────────────────────
    if data.experiences:
        section_heading("Work Experience")
        for exp in data.experiences:
            role    = exp.get("role", "")
            company = exp.get("company", "")
            start   = _fmt_date(exp.get("start_date"))
            end     = _fmt_date(exp.get("end_date"))

            row_p = doc.add_paragraph()
            row_p.paragraph_format.space_after = Pt(0)
            add_run(row_p, role, bold=True, color=DARK_RGB, size_pt=10.5)
            add_run(row_p, "  —  ", color=MID_RGB, size_pt=10)
            add_run(row_p, company, italic=True, color=ACCENT_RGB, size_pt=10)
            add_run(row_p, f"   {start} – {end}", color=LIGHT_RGB, size_pt=9)

            for b in _bullets(exp.get("description"), 6):
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.left_indent  = Inches(0.2)
                bp.paragraph_format.space_after  = Pt(1)
                run = bp.add_run(b)
                run.font.size = Pt(9.5)
                run.font.color.rgb = MID_RGB

            doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # ── projects ────────────────────────────────────────────────────────
    if data.projects:
        section_heading("Projects")
        for proj in data.projects:
            name = proj.get("name", "")
            tech = proj.get("technologies", "")
            desc = proj.get("description", "")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            add_run(p, name, bold=True, color=DARK_RGB, size_pt=10)
            if tech:
                add_run(p, f"  |  {tech}", color=ACCENT_RGB, size_pt=9)
            if desc:
                dp = doc.add_paragraph(desc[:300])
                dp.paragraph_format.space_after = Pt(2)
                for run in dp.runs:
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = MID_RGB

    # ── education ───────────────────────────────────────────────────────
    if data.educations:
        section_heading("Education")
        for edu in data.educations:
            institution = edu.get("institution", "")
            degree      = edu.get("degree", "")
            field       = edu.get("field_of_study", "")
            start       = _fmt_date(edu.get("start_date"))
            end         = _fmt_date(edu.get("end_date"))
            degree_line = ", ".join(filter(None, [degree, field]))

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            add_run(p, institution, bold=True, color=DARK_RGB, size_pt=10.5)
            add_run(p, f"   {start} – {end}", color=LIGHT_RGB, size_pt=9)
            if degree_line:
                dp = doc.add_paragraph(degree_line)
                dp.paragraph_format.space_after = Pt(3)
                for run in dp.runs:
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = MID_RGB

    # ── certifications ───────────────────────────────────────────────────
    if data.certifications:
        section_heading("Certifications")
        for cert in data.certifications:
            name   = cert.get("name", "")
            issuer = cert.get("issuer", "")
            issued = _fmt_date(cert.get("issue_date"))
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            add_run(p, name, bold=True, color=DARK_RGB, size_pt=10)
            if issuer:
                add_run(p, f"  —  {issuer}", color=MID_RGB, size_pt=9.5)
            if cert.get("issue_date"):
                add_run(p, f"  ({issued})", color=LIGHT_RGB, size_pt=9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Public API ─────────────────────────────────────────────────────────────

class ResumeExporter:
    def __init__(self, db: Session, user: User, app_id: int | None = None):
        self.db     = db
        self.user   = user
        self.app_id = app_id

    def _load(self) -> _ResumeData:
        return _ResumeData(self.db, self.user.id, self.user.email, app_id=self.app_id)

    def as_pdf(self) -> bytes:
        try:
            return _build_pdf(self._load())
        except Exception as exc:
            logger.error("resume_pdf_error", error=str(exc))
            raise HTTPException(status_code=500, detail=f"PDF generation failed: {exc}")

    def as_docx(self) -> bytes:
        try:
            return _build_docx(self._load())
        except Exception as exc:
            logger.error("resume_docx_error", error=str(exc))
            raise HTTPException(status_code=500, detail=f"DOCX generation failed: {exc}")
